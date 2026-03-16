import os
import io
import json
import pandas as pd
from pypdf import PdfReader
from docx import Document
from pinecone import Pinecone
from groq import Groq
from sentence_transformers import SentenceTransformer
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
from dotenv import load_dotenv

# Load keys from .env file — works whether run directly or imported by api.py
_env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", ".env")
load_dotenv(dotenv_path=os.path.abspath(_env_path))
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
FOLDER_ID = os.getenv("FOLDER_ID")

if not all([PINECONE_API_KEY, GROQ_API_KEY, FOLDER_ID]):
    missing = [k for k, v in {"PINECONE_API_KEY": PINECONE_API_KEY, "GROQ_API_KEY": GROQ_API_KEY, "FOLDER_ID": FOLDER_ID}.items() if not v]
    raise ValueError(f"Missing environment variables: {', '.join(missing)}")

import logging

# Set paths relative to script location
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
REPO_DIR = os.path.dirname(SCRIPT_DIR)

# Full drive scope so we can read AND upload files from the frontend
SCOPES = ["https://www.googleapis.com/auth/drive"]
UPLOADED_FILES_RECORD = os.path.join(REPO_DIR, "uploaded_files.json")  # Local tracker file

# Connect
pc = Pinecone(api_key=PINECONE_API_KEY)
groq_client = Groq(api_key=GROQ_API_KEY)
index = pc.Index("employee-rag")
model = SentenceTransformer("all-MiniLM-L6-v2")

# ─────────────────────────────────────────
# SESSION MEMORY
# ─────────────────────────────────────────
# Stores chat history: { session_id: [ {"role": "user", "content": "..."}, {"role": "assistant", "content": "..."} ] }
SESSION_MEMORY = {}

# ─────────────────────────────────────────
# AGENT CONFIGURATIONS
# ─────────────────────────────────────────
AGENT_CONFIGS = {
    "interview_specialist": {
        "name": "Interview Specialist",
        "system_prompt": "You are a professional HR Interview Specialist. Your expertise is in analyzing interview question datasets and providing detailed guidance on interview techniques and expected answers. Answer questions thoroughly based on the provided data.",
    },
    "resume_analyst": {
        "name": "Resume Analyst",
        "system_prompt": "You are an expert Resume Analyst. Your role is to scrutinize candidate resumes, highlight key skills, identify experience gaps, and answer specific questions about individual candidates' backgrounds based on the provided resumes.",
    },
    "screening_report_analyzer": {
        "name": "Screening Report Analyzer",
        "system_prompt": "You are a detailed Screening Report Analyzer. Your task is to review post-interview screening reports from employers and provide comprehensive summaries and insights into candidate performance and employer feedback.",
    }
}

# ─────────────────────────────────────────
# UPLOAD TRACKER
# ─────────────────────────────────────────
def load_uploaded_record():
    """Load the set of already-uploaded file IDs from local JSON."""
    if os.path.exists(UPLOADED_FILES_RECORD):
        with open(UPLOADED_FILES_RECORD, "r") as f:
            return set(json.load(f))
    return set()

def save_uploaded_record(uploaded_ids: set):
    """Persist the updated set of uploaded file IDs."""
    with open(UPLOADED_FILES_RECORD, "w") as f:
        json.dump(list(uploaded_ids), f)

# ─────────────────────────────────────────
# GOOGLE DRIVE
# ─────────────────────────────────────────
def get_drive_service():
    creds = None
    token_path = os.path.join(REPO_DIR, "token.json")
    credentials_path = os.path.join(REPO_DIR, "credentials.json")
    
    if os.path.exists(token_path):
        creds = Credentials.from_authorized_user_file(token_path, SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(credentials_path, SCOPES)
            creds = flow.run_local_server(port=0)
        with open(token_path, "w") as token:
            token.write(creds.to_json())
    return build("drive", "v3", credentials=creds)

def download_file(service, file_id, filename, mime_type=None):
    if mime_type and mime_type.startswith("application/vnd.google-apps."):
        # Export Google Docs/Sheets/Slides to PDF for extraction
        request = service.files().export_media(fileId=file_id, mimeType="application/pdf")
    else:
        request = service.files().get_media(fileId=file_id)
    
    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    buf.seek(0)
    return buf

def upload_file_to_drive(file_bytes: bytes, filename: str) -> str:
    """Upload a file to the configured Google Drive folder. Returns the Drive file ID."""
    service = get_drive_service()
    mime_map = {
        ".pdf": "application/pdf",
        ".csv": "text/csv",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    ext = os.path.splitext(filename)[1].lower()
    mime_type = mime_map.get(ext, "application/octet-stream")

    file_metadata = {"name": filename, "parents": [FOLDER_ID]}
    media = MediaIoBaseUpload(io.BytesIO(file_bytes), mimetype=mime_type, resumable=True)
    uploaded = service.files().create(
        body=file_metadata,
        media_body=media,
        fields="id"
    ).execute()
    return uploaded.get("id")

def chunk_text(text: str, max_size: int = 500) -> list[str]:
    """Splits text into chunks of maximum max_size characters, preserving words."""
    chunks = []
    raw_text = str(text)
    # Replace newlines with spaces for smoother chunking
    processed_text = " ".join(raw_text.split())
    
    while len(processed_text) > 0:
        if len(processed_text) <= max_size:
            chunks.append(processed_text)
            break
        
        # Find last space within max_size
        split_at = processed_text.rfind(' ', 0, max_size)
        if split_at == -1:
            split_at = max_size
        
        chunk = processed_text[:split_at].strip()
        if chunk:
            chunks.append(chunk)
        processed_text = processed_text[split_at:].strip()
    return chunks

def extract_text(buf, filename, chunk_size=500, mime_type=None):
    texts = []
    raw_texts = []
    filename_lower = filename.lower()
    
    if filename_lower.endswith(".csv") or mime_type == "text/csv":
        try:
            buf.seek(0)
            df = pd.read_csv(buf)
        except Exception:
            # Try alternate encoding for legacy CSVs
            buf.seek(0)
            df = pd.read_csv(buf, encoding="latin-1")
            
        for _, row in df.iterrows():
            text = " | ".join([f"{col}: {val}" for col, val in row.items()])
            raw_texts.append(text)
    elif filename_lower.endswith(".pdf") or mime_type in ["application/pdf", "application/vnd.google-apps.document", "application/vnd.google-apps.spreadsheet", "application/vnd.google-apps.presentation"]:
        # Note: Google-apps types are exported to PDF in download_file()
        try:
            reader = PdfReader(buf)
            for page in reader.pages:
                text = page.extract_text()
                if text and text.strip():
                    raw_texts.append(text)
        except Exception as e:
            print(f"  ❌ PDF extraction failed for {filename}: {e}")
    elif filename_lower.endswith(".docx") or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(buf)
        for para in doc.paragraphs:
            if para.text.strip():
                raw_texts.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    raw_texts.append(row_text)
    elif filename_lower.endswith(".txt") or mime_type == "text/plain":
        buf.seek(0)
        text = buf.read().decode("utf-8", errors="replace")
        if text.strip():
            raw_texts.append(text)
                
    for rt in raw_texts:
        chunks = chunk_text(rt, max_size=chunk_size)
        texts.extend(chunks)
        
    return texts

def sync_new_files_from_drive():
    """
    Checks Google Drive for files not yet uploaded to Pinecone.
    Returns extracted text chunks only from new files.
    """
    print("🔗 Connecting to Google Drive...")
    service = get_drive_service()

    all_drive_files = []
    page_token = None
    while True:
        results = service.files().list(
            q=f"'{FOLDER_ID}' in parents and mimeType != 'application/vnd.google-apps.folder'",
            fields="nextPageToken, files(id, name, mimeType, size)",
            pageSize=1000,
            pageToken=page_token
        ).execute()
        all_drive_files.extend(results.get("files", []))
        page_token = results.get("nextPageToken")
        if not page_token:
            break
            
    # SMART SORT: Prioritize PDFs/DOCXs over CSVs and smaller files over larger ones
    def sort_priority(f):
        m = f.get('mimeType', '').lower()
        size = int(f.get('size', 0))
        priority = 100
        if "pdf" in m or "word" in m or "document" in m:
            priority = 0
        elif "csv" in m or "sheet" in m:
            priority = 10
        return (priority, size)
    
    all_drive_files.sort(key=sort_priority)
    print(f"📁 Found {len(all_drive_files)} total files in Drive folder (Sorted by priority)")

    uploaded_ids = load_uploaded_record()
    new_files = [f for f in all_drive_files if f["id"] not in uploaded_ids]

    if not new_files:
        print("✅ No new files to upload. Pinecone is already up to date.")
        return [], uploaded_ids

    print(f"🆕 {len(new_files)} new file(s) detected — processing...")
    all_texts = []
    newly_uploaded_ids = set()

    for file in new_files:
        filename = file["name"]
        file_id = file["id"]
        file_size = int(file.get("size", 0))
        if file_size > 5 * 1024 * 1024: # 5MB limit for priority sync
            print(f"  ⏭️ Skipping large file {filename} ({file_size / 1024 / 1024:.2f}MB) for now...")
            continue
            
        print(f"  Reading {filename}...")
        try:
            buf = download_file(service, file_id, filename, mime_type=file.get("mimeType"))
            # If it was a Google Doc exported to PDF, pretend it's a PDF for extraction
            mime_type = file.get("mimeType")
            texts = extract_text(buf, filename, mime_type=mime_type)
            if texts:
                # Pass current filename to store_in_pinecone for metadata tagging
                store_in_pinecone(texts, {file_id}, filename=filename)
                all_texts.extend(texts)
                newly_uploaded_ids.add(file_id)
                print(f"  ✅ Extracted and indexed {len(texts)} chunks from {filename}")
            else:
                print(f"  ⚠️  No text extracted from {filename} — skipping")
        except Exception as e:
            print(f"  ❌ Could not read {filename}: {e}")

    updated_ids = set(uploaded_ids).union(set(newly_uploaded_ids))
    return all_texts, updated_ids

# ─────────────────────────────────────────
# PINECONE
# ─────────────────────────────────────────
def store_in_pinecone(texts, updated_ids, filename="unknown"):
    print(f"\n📤 Uploading to Pinecone (Source: {filename})...")
    texts = [t for t in texts if len(t.strip()) >= 10]
    if not texts:
        return

    # Use current index stats to offset new chunk IDs and avoid overwriting old ones
    stats = index.describe_index_stats()
    existing_count = stats.get("total_vector_count", 0)

    batch_size = 100
    for i in range(0, len(texts), batch_size):
        batch_texts = texts[i:i + batch_size]
        encoded = model.encode(batch_texts)
        embeddings = encoded.tolist()
        vectors = []
        for j, emb in enumerate(embeddings):
            # Deterministic ID based on filename and chunk index
            clean_filename = "".join(x for x in filename if x.isalnum() or x in "._-")
            chunk_id = f"{clean_filename}_chunk_{i + j}"
            
            # Smart context: prepend source info to the text itself for better retrieval
            # This ensures terms like "Atul Mathur" in the filename help match the chunk
            full_text = f"NAME/SOURCE: {filename} | CONTENT: {batch_texts[j]}"
            
            vectors.append({
                "id": chunk_id,
                "values": emb,
                "metadata": {
                    "text": full_text,
                    "source": filename
                }
            })
        index.upsert(vectors=vectors)
        print(f"  Uploaded {min(i + batch_size, len(texts))}/{len(texts)} chunks...")

    # Load and merge with existing record before saving
    all_known_ids = load_uploaded_record()
    all_known_ids.update(updated_ids)
    save_uploaded_record(all_known_ids)
    print(f"✅ Data from '{filename}' indexed in Pinecone!")

# ─────────────────────────────────────────
# ASK QUESTIONS
# ─────────────────────────────────────────
def ask_question(question, agent_type="interview_specialist", session_id="default"):
    # Get configuration for the selected agent
    config = AGENT_CONFIGS.get(agent_type, AGENT_CONFIGS["interview_specialist"])
    system_prompt = config["system_prompt"]

    # Retrieval from Pinecone
    question_embedding = model.encode(question).tolist()
    results = index.query(
        vector=question_embedding,
        top_k=20,  # Increased top_k for better candidate coverage
        include_metadata=True
    )
    
    # Sort matches by score for quality (query already does this, but for clarity)
    matches = sorted(results["matches"], key=lambda x: x["score"], reverse=True)
    
    # Format context with source labels
    context_parts = []
    for r in matches:
        source = r["metadata"].get("source", "unknown")
        chunk_text = r["metadata"].get("text", "")
        if chunk_text:
            context_parts.append(f"--- DOCUMENT: {source} ---\n{chunk_text}")
    
    context = "\n\n".join(context_parts)
    
    # Construct complete system instructions
    instruction = f"""{system_prompt}

You are an expert HR assistant. Base your answer ONLY on the provided Context Data.
If the information is not present, state that clearly.

Context Data:
{context}"""

    # Manage session memory (keep last 10 messages for context)
    if session_id not in SESSION_MEMORY:
        SESSION_MEMORY[session_id] = []
    
    # Build messages for the LLM
    messages = [{"role": "system", "content": instruction}]
    
    # Append history
    messages.extend(SESSION_MEMORY[session_id][-10:])
    
    # Append current question
    messages.append({"role": "user", "content": question})

    response = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=messages,
        temperature=0.2
    )
    
    answer = response.choices[0].message.content
    
    # Update memory
    SESSION_MEMORY[session_id].append({"role": "user", "content": question})
    SESSION_MEMORY[session_id].append({"role": "assistant", "content": answer})
    
    return answer

# ─────────────────────────────────────────
# MAIN — only runs when executed directly
# ─────────────────────────────────────────
if __name__ == "__main__":
    new_texts, updated_ids = sync_new_files_from_drive()

    if new_texts:
        store_in_pinecone(new_texts, updated_ids)
    else:
        print("⏭️  Skipping Pinecone upload — nothing new to store.")

    print("\n🤖 Assistant ready! Ask anything about your data (type 'quit' to exit)\n")
    while True:
        question = input("You: ")
        if question.lower() == "quit":
            break
        answer = ask_question(question)
        print(f"\nAssistant: {answer}\n")